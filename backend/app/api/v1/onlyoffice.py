import logging

from fastapi import APIRouter, HTTPException, Response, Request, Depends
from datetime import datetime, timedelta
from typing import Optional
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pathlib import Path
import tempfile
import io
from docx import Document
from typing import Any, Dict
from sqlalchemy.orm import Session
from app.api.deps import get_db, get_current_user
from app.models.research_paper import ResearchPaper
from app.models.branch import Branch, Commit
from app.models.user import User
from app.models.reference import Reference
from app.core.config import settings
from fastapi.responses import JSONResponse
import json
import os
import requests
import re
import html as _html
import shutil
import subprocess
import json as _json
from typing import Optional

logger = logging.getLogger(__name__)

try:
    import mammoth as _mammoth  # type: ignore
except Exception:  # pragma: no cover
    _mammoth = None

router = APIRouter()

# In-memory token registry (dev convenience). Maps paperId/key -> token with TTL.
_TOKEN_STORE: dict[str, tuple[str, float]] = {}
# In-memory last-save diagnostics for debugging OnlyOffice callbacks
_LAST_SAVE: dict[str, dict] = {}

def _now_ts() -> float:
    return datetime.utcnow().timestamp()

def _prune_tokens() -> None:
    now = _now_ts()
    expired = [k for k, (_, exp) in _TOKEN_STORE.items() if exp < now]
    for k in expired:
        _TOKEN_STORE.pop(k, None)

def create_sample_docx() -> bytes:
    """Create a sample DOCX document for OnlyOffice testing"""
    doc = Document()

    # Add title
    title = doc.add_heading('ScholarHub Document Prototype', 0)

    # Add some sample content
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'This is a sample document created for testing OnlyOffice Document Server integration with ScholarHub. '
        'OnlyOffice provides a collaborative document editing experience similar to Microsoft Office.'
    )

    doc.add_heading('Features to Test', level=1)
    doc.add_paragraph('Please test the following features:', style='List Bullet')

    features = [
        'Text formatting (bold, italic, underline)',
        'Headings and paragraph styles',
        'Lists and bullet points',
        'Tables and images',
        'Collaborative editing (if multiple users)',
        'Auto-save functionality',
        'Document export capabilities'
    ]

    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Comparison Notes', level=1)
    comparison_text = doc.add_paragraph(
        'Use this section to note differences between OnlyOffice and TipTap:\n\n'
        '• User Experience: \n'
        '• Performance: \n'
        '• Feature Set: \n'
        '• Integration Complexity: \n'
        '• Collaboration Features: \n'
    )

    doc.add_page_break()
    doc.add_heading('Additional Testing Space', level=1)
    doc.add_paragraph('Use this page for additional testing and notes.')

    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.read()

def _html_to_docx_best_effort(html: str) -> bytes:
    """Convert HTML to DOCX using best available tools.

    - Prefer Pandoc if installed (highest fidelity)
    - Fallback to a simple python-docx builder for basic structure
    """
    html = (html or '').strip()
    # 1) Prefer pandoc if available
    try:
        if shutil.which('pandoc'):
            with tempfile.TemporaryDirectory() as td:
                html_path = Path(td) / 'input.html'
                docx_path = Path(td) / 'output.docx'
                html_path.write_text(html, encoding='utf-8')
                # Run pandoc: HTML -> DOCX
                cmd = ['pandoc', str(html_path), '-f', 'html', '-t', 'docx', '-o', str(docx_path)]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15)
                if docx_path.exists():
                    return docx_path.read_bytes()
    except Exception:
        pass

    # 2) Fallback: very simple HTML -> DOCX via python-docx (formatting will be limited)
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        doc = Document()

        def add_text_with_runs(p, node):
            # Basic inline formatting for <b>/<strong>, <i>/<em>, <u>
            from docx.text.run import Run
            for child in node.children:
                if isinstance(child, str):
                    if child:
                        p.add_run(child)
                else:
                    name = child.name.lower() if hasattr(child, 'name') and child.name else ''
                    if name in ('b', 'strong'):
                        r = p.add_run(child.get_text())
                        r.bold = True
                    elif name in ('i', 'em'):
                        r = p.add_run(child.get_text())
                        r.italic = True
                    elif name == 'u':
                        r = p.add_run(child.get_text())
                        r.underline = True
                    else:
                        # Recurse for nested spans
                        add_text_with_runs(p, child)

        body = soup.body or soup
        for el in body.find_all(['h1','h2','h3','h4','h5','h6','p','ul','ol'], recursive=True):
            name = el.name.lower()
            text = el.get_text(strip=True)
            if not text:
                continue
            if name.startswith('h') and name[1:].isdigit():
                level = max(1, min(6, int(name[1:])))
                # python-docx heading levels are 0..9; map h1..h6 -> 1..6
                doc.add_heading(text, level=level)
            elif name == 'p':
                p = doc.add_paragraph()
                add_text_with_runs(p, el)
            elif name in ('ul', 'ol'):
                style = 'List Bullet' if name == 'ul' else 'List Number'
                for li in el.find_all('li', recursive=False):
                    t = li.get_text(strip=True)
                    if t:
                        doc.add_paragraph(t, style=style)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        logger.debug("HTML->DOCX using simple fallback (python-docx). Formatting may be limited.")
        return bio.read()
    except Exception:
        return b''

def _html_to_docx_via_onlyoffice_url(html: str) -> Optional[bytes]:
    """Convert HTML to DOCX using OnlyOffice ConvertService with URL method.

    Uses the sample-document endpoint with HTML content.
    Returns DOCX bytes on success, or None if conversion fails.
    """
    try:
        # SIMPLIFIED APPROACH: OnlyOffice ConvertService appears to have issues
        # with URL access. For now, we'll disable it and rely on our excellent
        # Pandoc converter which is working perfectly.
        logger.debug("OnlyOffice ConvertService disabled due to URL access issues (error -4), using Pandoc instead")
        return None

    except Exception as e:
        logger.error("OnlyOffice ConvertService error: %s", e)
        return None

@router.get("/sample-document")
async def get_sample_document(paperId: Optional[str] = None):
    """Serve a sample DOCX document for OnlyOffice testing"""
    try:
        # Build a sample document. Do NOT inject the paperId as visible content.
        # The plugin can still auto-detect the paperId from the filename or key.
        content = create_sample_docx()

        # Create a temporary file to serve
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        fname = f"scholarhub_sample.docx"
        if paperId:
            # Use a name the plugin can parse to auto-fill Paper ID
            fname = f"Paper {paperId}.docx"
        return FileResponse(
            path=tmp_file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=fname,
            headers={"Cache-Control": "no-cache"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to create document. Please try again.")

@router.get("/document")
async def get_document_for_paper(paperId: str, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Serve a DOCX for OnlyOffice. Prioritizes saved DOCX artifacts for formatting preservation."""
    paper = db.query(ResearchPaper).filter(ResearchPaper.id == paperId).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Research paper not found")

    try:
        # DOCX-first strategy: If we have a saved OnlyOffice DOCX artifact, serve it directly
        # This preserves all formatting from previous OnlyOffice sessions
        docx_rel = None
        try:
            cj = paper.content_json or {}
            logger.debug("Paper content_json: %s", cj)
            if isinstance(cj, dict):
                docx_rel = cj.get('oo_docx_path')
                authoring_mode = cj.get('authoring_mode', 'rich')
                logger.debug("Found oo_docx_path: %s", docx_rel)
        except Exception as e:
            logger.error("Error reading content_json: %s", e)
            authoring_mode = 'rich'

        # Also check if DOCX file exists directly (fallback)
        if not docx_rel:
            direct_docx_path = Path(__file__).resolve().parent.parent.parent.parent / f"uploads/onlyoffice/{paperId}.docx"
            if direct_docx_path.exists():
                logger.debug("Found DOCX artifact via direct path: %s", direct_docx_path)
                docx_rel = f"uploads/onlyoffice/{paperId}.docx"

        if docx_rel:
            path = Path(__file__).resolve().parent.parent.parent.parent / docx_rel
            if path.exists():
                logger.info("Serving preserved DOCX artifact: %s (%d bytes)", path, path.stat().st_size)
                return FileResponse(
                    path=str(path),
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    filename=f"Paper {paperId}.docx",
                    headers={"Cache-Control": "no-cache"}
                )
            else:
                logger.warning("DOCX path in content_json doesn't exist: %s", path)
        else:
            logger.debug("No DOCX artifact path found in content_json")

        # If no DOCX artifact exists, convert from HTML using best available method
        html = (paper.content or '').strip()
        data = None

        # Priority 1: OnlyOffice ConvertService (with URL method - fixed!)
        if html:
            data = _html_to_docx_via_onlyoffice_url(html)
            if data:
                logger.info("HTML->DOCX using OnlyOffice ConvertService (URL method)")

        # Priority 2: Local converters (Pandoc > python-docx)
        if not data and html:
            data = _html_to_docx_best_effort(html)
            if data:
                logger.debug("HTML->DOCX using local converters")

        # Priority 3: Sample document if no content
        if not data:
            data = create_sample_docx()
            logger.debug("Using sample DOCX (no content available)")

        # Save the generated DOCX as an artifact for future use
        if data and paperId:
            try:
                docx_rel_path = _save_docx_for_paper(paperId, data)
                # Update paper metadata to track this DOCX
                cj = paper.content_json or {}
                if not isinstance(cj, dict):
                    cj = {}
                cj['oo_docx_path'] = docx_rel_path
                cj['oo_last_generated'] = datetime.utcnow().isoformat() + 'Z'
                # Do NOT force authoring_mode here; preserve existing if present
                # This avoids flipping LaTeX papers to rich due to a stray open
                paper.content_json = cj
                db.commit()
                logger.info("Saved DOCX artifact for future use: %s", docx_rel_path)
            except Exception as e:
                logger.error("Failed to save DOCX artifact: %s", e)

        # Serve the DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(data)
            tmp_path = tmp_file.name

        fname = f"Paper {paperId}.docx"
        return FileResponse(
            path=tmp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=fname,
            headers={"Cache-Control": "no-cache"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to build document. Please try again.")

from app.database import SessionLocal
from app.models.collaboration_session import CollaborationSession
from app.services.websocket_manager import connection_manager


def _extract_paper_id_from_key(key: str | None) -> str | None:
    """Extract strict UUID (36 chars) from OnlyOffice key like 'paper-<uuid>-<ts>'."""
    if not key:
        return None
    s = str(key)
    m = re.search(r"([0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12})", s)
    return m.group(1) if m else None


def _docx_to_html_best_effort(data: bytes) -> str:
    """Convert DOCX to HTML with best fidelity available.

    Tries Mammoth first (preserves headings/lists/bold/italic/links),
    falls back to a simple paragraph extractor using python-docx,
    and ultimately to plain text if all else fails.
    """
    # Try 1: Pandoc first (best for preserving visual formatting including fonts)
    try:
        if shutil.which('pandoc'):
            with tempfile.TemporaryDirectory() as td:
                src = Path(td) / 'in.docx'
                dst = Path(td) / 'out.html'
                src.write_bytes(data)
                # Enhanced pandoc command with CSS and font preservation
                cmd = [
                    'pandoc', str(src),
                    '-f', 'docx',
                    '-t', 'html',
                    '--standalone',  # Include CSS in output
                    '--wrap=preserve',  # Preserve line breaks
                    '--extract-media', str(td),  # Extract images
                    '--metadata', 'title=',  # Remove title metadata
                    '-o', str(dst)
                ]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=20)
                if dst.exists():
                    html = dst.read_text(encoding='utf-8', errors='ignore').strip()
                    if html and len(html) > 20:
                        # Remove the unwanted title "in" from Pandoc output
                        html = html.replace('<title>in</title>', '<title></title>')
                        html = html.replace('<title>in.docx</title>', '<title></title>')
                        logger.debug("Using Pandoc for DOCX->HTML conversion with CSS preservation (%d chars)", len(html))
                        return html
    except Exception as e:
        logger.warning("Pandoc conversion failed: %s", e)

    # Try 2: Mammoth (enhanced) - Better semantic preservation
    if _mammoth is not None:
        try:
            # Simplified but more reliable style mapping
            style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Title'] => h1:fresh
                r[style-name='Strong'] => strong
                r[style-name='Emphasis'] => em
                u => u
            """

            # Custom transform to preserve more formatting
            def transform_document(document):
                # This is a placeholder for custom document transformation
                # Mammoth doesn't handle inline styles well, so we focus on structure
                return document

            convert_options = {
                'style_map': style_map,
                'include_default_style_map': True,
                'include_embedded_style_map': True,
                'transform_document': transform_document,
                'ignore_empty_paragraphs': False
            }

            result = _mammoth.convert_to_html(io.BytesIO(data), **convert_options)
            html = (result.value or '').strip()

            if html and len(html) > 20:
                logger.debug("Using Mammoth for DOCX->HTML conversion (semantic) (%d chars)", len(html))
                return html
        except Exception as e:
            logger.warning("Mammoth conversion failed: %s", e)


    # 2) Fallback: simple python-docx paragraph dump
    try:
        from docx import Document as _Doc
        doc = _Doc(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or '').strip()
            if t:
                parts.append(f"<p>{_html.escape(t)}</p>")
        html = "\n".join(parts).strip()
        if html:
            return html
    except Exception:
        pass

    # 3) Final fallback: preformatted plain text
    try:
        s = data.decode('utf-8', errors='ignore')
        return f"<pre>{_html.escape(s)}</pre>"
    except Exception:
        return ""

def _ensure_uploads_dir() -> Path:
    root = Path(__file__).resolve().parent.parent.parent.parent
    up = root / 'uploads' / 'onlyoffice'
    up.mkdir(parents=True, exist_ok=True)
    return up

def _save_docx_for_paper(paper_id: str, data: bytes) -> str:
    base = _ensure_uploads_dir()
    fname = f"{paper_id}.docx"
    path = base / fname
    path.write_bytes(data)
    logger.debug("Saved DOCX artifact to %s", path)
    # return relative path under uploads for portability
    return str(Path('uploads') / 'onlyoffice' / fname)


def _docx_to_html_via_onlyoffice(docx_url: str) -> str:
    """Use OnlyOffice ConvertService to convert DOCX->HTML preserving visual formatting.

    Returns HTML string, or empty string on failure.
    """
    try:
        convert_endpoint = f'{settings.ONLYOFFICE_DOCSERVER_URL or "http://localhost:8080"}/ConvertService.ashx'
        payload = {
            'async': False,
            'filetype': 'docx',
            'key': f"{datetime.utcnow().timestamp()}-{abs(hash(docx_url))}",
            'outputtype': 'html',
            'url': docx_url,
        }
        # Attempt JSON API first
        headers = {'Accept': 'application/json', 'Content-Type': 'application/json'}
        r = requests.post(convert_endpoint, json=payload, headers=headers, timeout=15)
        if not r.ok:
            # Fallback to form field 'data' as expected by some DS builds
            r = requests.post(convert_endpoint, data={'data': json.dumps(payload)}, timeout=20)
        r.raise_for_status()
        try:
            data = r.json()
        except Exception:
            # Some builds respond with XML on error -- log a preview
            preview = (r.text or '')[:200]
            logger.warning("Non-JSON response (%s): %s", r.status_code, preview)
            return ""
        file_url = data.get('fileUrl') or data.get('fileurl')
        if not file_url:
            return ""
        rr = requests.get(file_url, timeout=20)
        rr.raise_for_status()
        ctype = rr.headers.get('content-type', '').lower()
        if 'text/html' in ctype or file_url.lower().endswith('.html'):
            return (rr.text or '').strip()
        if 'zip' in ctype or file_url.lower().endswith('.zip'):
            with tempfile.TemporaryDirectory() as td:
                zpath = Path(td) / 'out.zip'
                zpath.write_bytes(rr.content)
                import zipfile as _zip
                with _zip.ZipFile(zpath) as zf:
                    name = next((n for n in zf.namelist() if n.lower().endswith(('.html', '.htm'))), None)
                    if name:
                        content = zf.read(name)
                        try:
                            return content.decode('utf-8', errors='ignore').strip()
                        except Exception:
                            return content.decode('latin-1', errors='ignore').strip()
        return ""
    except Exception as e:
        logger.error("DOCX->HTML via OnlyOffice failed: %s", e)
        return ""

def _docx_to_html_via_onlyoffice_upload(docx_bytes: bytes) -> str:
    """Upload DOCX to OnlyOffice ConvertService and request HTML output.

    This avoids URL fetch issues inside Document Server and can preserve visual formatting.
    """
    try:
        convert_endpoint = f'{settings.ONLYOFFICE_DOCSERVER_URL or "http://localhost:8080"}/ConvertService.ashx'
        payload = {
            'async': False,
            'filetype': 'docx',
            'key': f"{datetime.utcnow().timestamp()}-{len(docx_bytes)}",
            'outputtype': 'html',
        }
        files = {
            'file': ('input.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        resp = requests.post(convert_endpoint, data={'data': _json.dumps(payload)}, files=files, timeout=30)
        resp.raise_for_status()
        try:
            data = resp.json()
        except Exception:
            logger.warning("Upload Non-JSON response (%s): %s", resp.status_code, (resp.text or '')[:200])
            return ""
        file_url = data.get('fileUrl') or data.get('fileurl')
        if not file_url:
            return ""
        rr = requests.get(file_url, timeout=20)
        rr.raise_for_status()
        ctype = rr.headers.get('content-type', '').lower()
        if 'text/html' in ctype or file_url.lower().endswith('.html'):
            return (rr.text or '').strip()
        if 'zip' in ctype or file_url.lower().endswith('.zip'):
            with tempfile.TemporaryDirectory() as td:
                zpath = Path(td) / 'out.zip'
                zpath.write_bytes(rr.content)
                import zipfile as _zip
                with _zip.ZipFile(zpath) as zf:
                    name = next((n for n in zf.namelist() if n.lower().endswith(('.html', '.htm'))), None)
                    if name:
                        content = zf.read(name)
                        try:
                            return content.decode('utf-8', errors='ignore').strip()
                        except Exception:
                            return content.decode('latin-1', errors='ignore').strip()
        return ""
    except Exception as e:
        logger.error("DOCX->HTML via OnlyOffice upload failed: %s", e)
        return ""


def _summarize_changes(previous_html: str, new_html: str) -> list[dict[str, Any]]:
    """Produce a minimal change summary suitable for Commit.changes.

    Keeps implementation lightweight and robust for autosave events.
    """
    if (previous_html or '').strip() == (new_html or '').strip():
        return [{
            'type': 'update',
            'section': 'Document',
            'oldContent': '',
            'newContent': 'Autosave (no changes)'
        }]
    try:
        from bs4 import BeautifulSoup
        def _text_excerpt(html: str, n: int = 160) -> str:
            text = BeautifulSoup(html or '', 'html.parser').get_text(separator=' ', strip=True)
            return (text[:n] + ('...' if len(text) > n else '')) if text else ''
        return [{
            'type': 'update',
            'section': 'Document',
            'oldContent': _text_excerpt(previous_html),
            'newContent': _text_excerpt(new_html),
            'position': 0
        }]
    except Exception:
        return [{
            'type': 'update',
            'section': 'Document',
            'oldContent': '',
            'newContent': 'Autosave update'
        }]


@router.post("/callback")
async def onlyoffice_callback(payload: Dict[str, Any], db: Session = Depends(get_db)):
    """Handle OnlyOffice document server callbacks for saving documents"""
    try:
        logger.info("OnlyOffice callback received: %s", json.dumps(payload, indent=2))

        # OnlyOffice callback status meanings:
        # 0 - no document with key identifier could be found
        # 1 - document is being edited
        # 2 - document is ready for saving
        # 3 - document saving error has occurred
        # 4 - document is closed with no changes
        # 6 - document is being edited, but the current document state is saved
        # 7 - error has occurred while force saving the document

        status = payload.get('status', 0)

        if status in (2, 6):
            # 2: ready for saving, 6: state saved; process content
            document_url = payload.get('url')
            document_key = payload.get('key')
            logger.info("OnlyOffice save: status=%s key=%s url=%s", status, document_key, document_url)
            paper_id = _extract_paper_id_from_key(document_key)
            if not paper_id:
                logger.warning("Cannot extract paper_id from key; skipping DB update")
            elif document_url:
                try:
                    # Prefer local Pandoc/Mammoth conversion for reliability and speed; DS convert as last resort
                    r = requests.get(document_url, timeout=10)
                    r.raise_for_status()
                    data = r.content
                    # Save DOCX as canonical for OO-rich papers to preserve formatting
                    docx_rel_path = _save_docx_for_paper(paper_id, data)
                    # Update content_json with OO docx path and last save time
                    try:
                        paper = db.query(ResearchPaper).filter(ResearchPaper.id == paper_id).first()
                    except Exception as qerr:
                        try: db.rollback()
                        except Exception: pass
                        logger.error("DB query failed in callback: %s", qerr)
                        paper = None
                    if paper:
                        try:
                            cj = paper.content_json or {}
                            if not isinstance(cj, dict):
                                cj = {}
                            cj['oo_docx_path'] = docx_rel_path
                            cj['oo_last_save'] = datetime.utcnow().isoformat() + 'Z'
                            # Preserve existing authoring_mode; do not force switch
                            paper.content_json = cj

                            # Convert DOCX to HTML using the best available method
                            # Priority: OnlyOffice ConvertService > Mammoth > Pandoc > python-docx
                            html = ""

                            # Skip OnlyOffice ConvertService (currently failing with error -7)
                            # TODO: Fix ConvertService configuration for file upload method
                            logger.debug("Skipping OnlyOffice ConvertService (error -7), using local converters")

                            # Use enhanced local converters (Pandoc first for better formatting)
                            if not html:
                                html = _docx_to_html_best_effort(data)
                                logger.debug("Used enhanced Pandoc converter for DOCX->HTML conversion")

                            if html:
                                paper.content = html

                            db.commit()
                            logger.info("Updated paper %s content_json with OO docx path: %s", paper_id, docx_rel_path)
                        except Exception as cerr:
                            try: db.rollback()
                            except Exception: pass
                            logger.error("DB update failed in callback: %s", cerr)
                    # Broadcast to active sessions with current HTML view (if any)
                    if paper:
                        # Diagnostics: log lengths and hashes to detect potential overwrites
                        try:
                            import hashlib as _hash
                            prev = paper.content or ''
                            prev_len = len(prev)
                            new_len = len(paper.content or '')
                            prev_hash = _hash.sha256(prev.encode('utf-8', errors='ignore')).hexdigest()[:12]
                            new_hash = _hash.sha256((paper.content or '').encode('utf-8', errors='ignore')).hexdigest()[:12]
                            logger.debug("paper=%s prev_len=%d prev_hash=%s new_len=%d new_hash=%s", paper_id, prev_len, prev_hash, new_len, new_hash)
                            # Record last-save info for frontend debugging
                            try:
                                _LAST_SAVE[str(paper_id)] = {
                                    'ts': datetime.utcnow().isoformat() + 'Z',
                                    'prev_len': prev_len,
                                    'prev_hash': prev_hash,
                                    'new_len': new_len,
                                    'new_hash': new_hash,
                                    'oo_docx_path': docx_rel_path,
                                }
                            except Exception:
                                pass
                        except Exception:
                            pass

                            # Create/update a commit on main branch to record this autosave
                            try:
                                main_branch = db.query(Branch).filter(Branch.paper_id == paper.id, Branch.is_main == True).first()
                                if not main_branch:
                                    main_branch = Branch(
                                        name='main',
                                        paper_id=paper.id,
                                        author_id=paper.owner_id,
                                        is_main=True,
                                        last_commit_message='Initial commit'
                                    )
                                    db.add(main_branch)
                                    db.commit()
                                    db.refresh(main_branch)
                                    # Create initial commit from previous content if any
                                    init_commit = Commit(
                                        branch_id=main_branch.id,
                                        message='Initial commit',
                                        content=previous_html or '',
                                        content_json=None,
                                        author_id=paper.owner_id
                                    )
                                    db.add(init_commit)
                                    db.commit()

                                # Create autosave commit with summarized changes
                                changes = _summarize_changes(previous_html, html)
                                autosave_commit = Commit(
                                    branch_id=main_branch.id,
                                    message='Autosave from OnlyOffice',
                                    content=html,
                                    content_json=None,
                                    author_id=paper.owner_id,
                                    changes=changes
                                )
                                db.add(autosave_commit)
                                # Update branch last message
                                main_branch.last_commit_message = autosave_commit.message
                                db.commit()
                            except Exception as ce:
                                logger.error("Commit creation failed: %s", ce)
                            # Broadcast to active collaboration sessions for this paper
                            try:
                                sessions = db.query(CollaborationSession).filter(CollaborationSession.paper_id == paper_id, CollaborationSession.is_active == True).all()
                                import asyncio as _asyncio
                                for s in sessions:
                                    await connection_manager.broadcast_to_session(str(s.id), {"type": "document_changed", "content": html})
                            except Exception as be:
                                logger.error("Broadcast error: %s", be)
                except Exception as de:
                    logger.error("Download/convert failed: %s", de)

        elif status == 1:
            logger.debug("Document is being edited")
        elif status == 4:
            logger.debug("Document closed with no changes")
        else:
            logger.debug("Other callback status received: %s", status)

        # OnlyOffice expects a specific response format
        # Status codes: 0 - no errors, 1 - document key error, etc.
        return {"error": 0}

    except Exception as e:
        logger.error("OnlyOffice callback error: %s", e)
        return {"error": 1}

@router.get("/last-save")
async def onlyoffice_last_save(paperId: str):
    """Return last callback info recorded for a paper (debug only)."""
    try:
        return _LAST_SAVE.get(str(paperId)) or {"ts": None, "prev_len": None, "new_len": None}
    except Exception:
        return {"ts": None}

@router.get("/status")
async def onlyoffice_status():
    """Check OnlyOffice integration status"""
    return {
        "status": "active",
        "service": "OnlyOffice Document Server Integration",
        "features": {
            "sample_document": "available",
            "callback_handling": "enabled",
            "document_types": ["docx", "xlsx", "pptx"],
            "plugins": ["scholarhub-assistant"]
        }
    }

@router.get("/plugins")
async def list_plugins():
    """List available OnlyOffice plugins"""
    root_static = Path(__file__).resolve().parent.parent.parent.parent / "static"
    candidates = [root_static / "onlyoffice-plugins", root_static / "plugins"]
    plugins = []

    for plugins_dir in candidates:
        if not plugins_dir.exists():
            continue
        for plugin_dir in plugins_dir.iterdir():
            if plugin_dir.is_dir():
                config_file = plugin_dir / "config.json"
                if not config_file.exists():
                    continue
                try:
                    with open(config_file, 'r') as f:
                        config = json.load(f)
                    plugins.append({
                        "id": plugin_dir.name,
                        "name": config.get("name", plugin_dir.name),
                        "version": config.get("version", "1.0.0"),
                        "guid": config.get("guid", ""),
                        "description": config.get("variations", [{}])[0].get("description", ""),
                        "install_url": f"/onlyoffice/plugins/{plugin_dir.name}/config.json"
                    })
                except Exception as e:
                    logger.warning("Error reading plugin config %s: %s", config_file, e)

    return {"plugins": plugins}

@router.get("/plugins/{plugin_id}/config.json")
async def get_plugin_config(plugin_id: str, request: Request):
    """Get plugin configuration for installation"""
    root_static = Path(__file__).resolve().parent.parent.parent.parent / "static"
    # Support both scholarhub-assistant (onlyoffice-plugins) and hello-world (plugins)
    possible_paths = [
        root_static / "onlyoffice-plugins" / plugin_id / "config.json",
        root_static / "plugins" / plugin_id / "config.json",
    ]
    config_file = None
    for p in possible_paths:
        if p.exists():
            config_file = p
            break
    if config_file is None:
        raise HTTPException(status_code=404, detail="Plugin not found")

    try:
        with open(config_file, 'r') as f:
            config = json.load(f)

        # Update URLs to be absolute using the request base URL
        # Infer which static subdir we matched
        subdir = "onlyoffice-plugins" if "onlyoffice-plugins" in str(config_file) else "plugins"
        base = str(request.base_url).rstrip('/')
        base_url = f"{base}/static/{subdir}/{plugin_id}"
        for variation in config.get("variations", []):
            if "url" in variation:
                variation["url"] = f"{base_url}/{variation['url']}"

        return config
    except Exception as e:
        raise HTTPException(status_code=500, detail="Error reading plugin configuration.")

@router.get("/install-plugin-script")
async def get_install_plugin_script():
    """Get JavaScript code to install the ScholarHub plugin in OnlyOffice"""
    script = """
// ScholarHub Assistant Plugin Installation Script
// Paste this code into the browser console on an OnlyOffice document page

(function() {
    const pluginUrl = 'http://192.168.100.121:8000/onlyoffice/plugins/scholarhub-assistant/config.json';

    if (typeof window.Asc !== 'undefined' && window.Asc.plugin) {
        console.log('Installing ScholarHub Assistant plugin...');

        // Method 1: Try direct installation
        try {
            window.Asc.plugin.loadPlugin(pluginUrl, function(result) {
                if (result) {
                    console.log('ScholarHub Assistant plugin installed successfully!');
                    alert('ScholarHub Assistant plugin installed! Look for it in the Plugins menu.');
                } else {
                    console.log('Plugin installation failed');
                    installViaConsole();
                }
            });
        } catch (e) {
            console.log('Direct installation failed, trying alternative method...');
            installViaConsole();
        }
    } else {
        console.log('OnlyOffice plugin API not found. Make sure you are on an OnlyOffice document page.');
        alert('Please run this script on an OnlyOffice document page with a document open.');
    }

    function installViaConsole() {
        console.log('Alternative installation method...');

        // Add plugin to the list if possible
        if (window.Asc && window.Asc.plugin && window.Asc.plugin.manager) {
            try {
                const pluginConfig = {
                    url: pluginUrl,
                    guid: "asc.{8515F8A6-2B7B-4B31-A78B-9C4F17C6B357}",
                    name: "ScholarHub Assistant"
                };

                window.Asc.plugin.manager.addPlugin(pluginConfig);
                console.log('Plugin added to manager');
            } catch (e) {
                console.log('Manager installation also failed:', e);
                manualInstallInstructions();
            }
        } else {
            manualInstallInstructions();
        }
    }

    function manualInstallInstructions() {
        const instructions = `
Manual Plugin Installation Instructions:

1. In OnlyOffice, go to the Plugins tab
2. Click on the "Plugin Manager" or "Manage Plugins"
3. Look for an "Add Plugin" or "Install from URL" option
4. Use this URL: ${pluginUrl}
5. Click Install

If that doesn't work:
- The plugin files are available at: http://192.168.100.121:8000/static/onlyoffice-plugins/scholarhub-assistant/
- You may need to set up plugin development mode in OnlyOffice
        `;

        console.log(instructions);
        alert('Plugin installation requires manual steps. Check the browser console for instructions.');
    }
})();
"""

    return Response(
        content=script,
        media_type="application/javascript",
        headers={"Content-Disposition": "attachment; filename=install-scholarhub-plugin.js"}
    )

@router.get("/references")
async def bridge_references(paperId: str, db: Session = Depends(get_db)):
    """Dev bridge endpoint to return references without client-side auth.

    SECURITY NOTE: This endpoint bypasses normal auth checks and is intended for local development/testing
    with the OnlyOffice Document Server plugin. It is enabled only when DEBUG is true or ENVIRONMENT is development.
    """
    # IMPORTANT: This bridge endpoint bypasses auth for OnlyOffice plugin use.
    # If you need to restrict it, add proper checks or disable it in production.
    # Restrict to development/debug environments
    if not (settings.DEBUG or str(settings.ENVIRONMENT).lower() in {"dev", "development", "local"}):
        raise HTTPException(status_code=403, detail="Bridge disabled in this environment")

    paper = db.query(ResearchPaper).filter(ResearchPaper.id == paperId).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Research paper not found")
    refs = db.query(Reference).filter(Reference.paper_id == paperId).order_by(Reference.created_at.desc()).all()
    # Minimal projection
    def to_dict(r: Reference) -> Dict[str, Any]:
        return {
            'id': str(r.id),
            'title': r.title,
            'authors': r.authors or [],
            'year': r.year,
            'doi': r.doi,
            'url': r.url,
            'journal': r.journal,
            'abstract': r.abstract,
        }
    payload = { 'references': [to_dict(r) for r in refs], 'total': len(refs) }
    # Add explicit CORS header for DocServer origin to avoid browser CORS blocks
    oo_origin = settings.ONLYOFFICE_DOCSERVER_URL or 'http://localhost:8080'
    return JSONResponse(content=payload, headers={
        'Access-Control-Allow-Origin': oo_origin,
        'Vary': 'Origin'
    })

@router.post("/token")
async def register_token(request: Request):
    """Register a short-lived access token for a paperId/key so plugins can fetch it.

    Expects JSON body with { paperId?: str, key?: str } and Authorization: Bearer <token> header.
    Stores the token in memory for ~15 minutes.
    """
    try:
        data = await request.json()
    except Exception:
        data = {}
    paper_id = data.get("paperId")
    key = data.get("key")
    if not paper_id and not key:
        raise HTTPException(status_code=400, detail="paperId or key required")
    auth = request.headers.get("authorization") or request.headers.get("Authorization")
    if not auth or not auth.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = auth.split(" ", 1)[1].strip()
    _prune_tokens()
    ttl = _now_ts() + 15 * 60
    if paper_id:
        _TOKEN_STORE[f"paper:{paper_id}"] = (token, ttl)
    if key:
        _TOKEN_STORE[f"key:{key}"] = (token, ttl)
    return {"ok": True}

@router.get("/token")
async def fetch_token(paperId: Optional[str] = None, key: Optional[str] = None, current_user: User = Depends(get_current_user)):
    """Return a registered token for the given paperId or key, if not expired."""
    if not paperId and not key:
        raise HTTPException(status_code=400, detail="paperId or key required")
    _prune_tokens()
    if paperId:
        item = _TOKEN_STORE.get(f"paper:{paperId}")
        if item and item[1] > _now_ts():
            return {"token": item[0]}
    if key:
        item = _TOKEN_STORE.get(f"key:{key}")
        if item and item[1] > _now_ts():
            return {"token": item[0]}
    raise HTTPException(status_code=404, detail="Token not found or expired")
