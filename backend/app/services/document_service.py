import os
import uuid
import aiofiles
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import logging
from sqlalchemy.orm import Session
from app.models.document import Document, DocumentStatus, DocumentType
from app.models.document_chunk import DocumentChunk
from app.models.tag import Tag
from app.models.document_tag import DocumentTag
from app.schemas.document import DocumentUpload
from app.services.duplicate_detection_service import DuplicateDetectionService

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        self.duplicate_detector = DuplicateDetectionService()
        
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file to disk and return file path"""
        # Generate unique filename to prevent conflicts
        file_ext = Path(filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
            
        return str(file_path)
    
    def check_for_duplicates(self, db: Session, file_content: bytes, filename: str, 
                           extracted_text: str, owner_id: str) -> Dict[str, Any]:
        """Check for duplicate documents before saving"""
        return self.duplicate_detector.detect_all_duplicates(
            db, file_content, filename, extracted_text, owner_id
        )
    
    def detect_document_type(self, file_path: str, mime_type: str) -> DocumentType:
        """Detect document type based on file extension and MIME type"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf' or mime_type == 'application/pdf':
            return DocumentType.PDF
        elif file_ext in ['.docx', '.doc'] or mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            return DocumentType.DOCX
        elif file_ext == '.txt' or mime_type == 'text/plain':
            return DocumentType.TXT
        else:
            return DocumentType.UNKNOWN
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    def get_pdf_page_count(self, file_path: str) -> int:
        """Get the number of pages in a PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error getting PDF page count for {file_path}: {e}")
            return 0
    
    def get_docx_page_count(self, file_path: str) -> int:
        """Get the number of pages in a DOCX file (approximate)"""
        try:
            doc = DocxDocument(file_path)
            # DOCX doesn't have a direct page count, so we estimate based on content
            # This is a rough approximation - one page is roughly 500-800 words
            total_words = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            estimated_pages = max(1, total_words // 600)  # Assume 600 words per page
            return estimated_pages
        except Exception as e:
            logger.error(f"Error getting DOCX page count for {file_path}: {e}")
            return 0
    
    def get_txt_page_count(self, file_path: str) -> int:
        """Get the number of pages in a TXT file (approximate)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # Estimate pages based on character count (assuming ~2000 chars per page)
                estimated_pages = max(1, len(content) // 2000)
                return estimated_pages
        except Exception as e:
            logger.error(f"Error getting TXT page count for {file_path}: {e}")
            return 0
    
    def get_page_count(self, file_path: str, document_type: DocumentType) -> int:
        """Get page count based on document type"""
        if document_type == DocumentType.PDF:
            return self.get_pdf_page_count(file_path)
        elif document_type == DocumentType.DOCX:
            return self.get_docx_page_count(file_path)
        elif document_type == DocumentType.TXT:
            return self.get_txt_page_count(file_path)
        else:
            return 0
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text content from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str, document_type: DocumentType) -> str:
        """Extract text content based on document type"""
        if document_type == DocumentType.PDF:
            return self.extract_text_from_pdf(file_path)
        elif document_type == DocumentType.DOCX:
            return self.extract_text_from_docx(file_path)
        elif document_type == DocumentType.TXT:
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for AI processing"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # Try to find a good break point (end of sentence or paragraph)
            if end < len(text):
                # Look for sentence endings
                for i in range(end, min(end + 100, len(text))):
                    if text[i] in '.!?':
                        end = i + 1
                        break
                    elif text[i] == '\n':
                        end = i + 1
                        break
            
            chunks.append({
                'text': chunk_text.strip(),
                'start_char': start,
                'end_char': end,
                'chunk_metadata': {}
            })
            
            start = end - overlap
            if start >= len(text):
                break
                
        return chunks
    
    async def process_document(
        self, 
        db: Session, 
        document: Document, 
        file_content: bytes,
        tags: Optional[List[str]] = None
    ) -> Document:
        """Process uploaded document: extract text, chunk, and generate embeddings"""
        try:
            # Update status to processing
            document.status = DocumentStatus.PROCESSING
            db.commit()
            
            # Extract text from document
            text = self.extract_text(document.file_path, document.document_type)
            
            # Extract page count
            page_count = self.get_page_count(document.file_path, document.document_type)
            document.page_count = page_count
            
            # Chunk the text
            chunks = self.chunk_text(text)
            
            # Create document chunks in database
            for i, chunk_data in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_text=chunk_data['text'],
                    chunk_index=i,
                    chunk_metadata=chunk_data['chunk_metadata']
                )
                db.add(chunk)
            
            # Process tags if provided
            if tags:
                await self._process_tags(db, document, tags)
            
            # Update document status to processed
            document.status = DocumentStatus.PROCESSED
            db.commit()
            
            # Auto-process for AI
            try:
                logger.info(f"Auto-processing document {document.id} for AI...")
                from app.services.document_processing_service import DocumentProcessingService
                ai_processor = DocumentProcessingService()
                ai_success = ai_processor.process_document_for_ai(db, document)
                
                if ai_success:
                    logger.info(f"Document {document.id} successfully processed for AI")
                else:
                    logger.warning(f"Document {document.id} AI processing failed")
                    
            except Exception as ai_error:
                logger.error(f"Error auto-processing document {document.id} for AI: {ai_error}")
                # Don't fail the upload, just log the error
            
            logger.info(f"Document {document.id} processed successfully with {len(chunks)} chunks")
            return document
            
        except Exception as e:
            logger.error(f"Error processing document {document.id}: {e}")
            document.status = DocumentStatus.FAILED
            db.commit()
            raise
    
    async def _process_tags(self, db: Session, document: Document, tag_names: List[str]):
        """Process and create tags for a document"""
        for tag_name in tag_names:
            # Get or create tag
            tag = db.query(Tag).filter(Tag.name == tag_name).first()
            if not tag:
                tag = Tag(name=tag_name)
                db.add(tag)
                db.flush()  # Get the ID
            
            # Create document-tag relationship
            doc_tag = DocumentTag(document_id=document.id, tag_id=tag.id)
            db.add(doc_tag)
        
        db.commit()
